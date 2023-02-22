import openpyxl
import os
from docx import Document
import win32com.client


def banner():
    banner_text = '''
--------------------------------------------------------------------------------------------
               __            __        __    __   
  ____ ___  __/ /_____      / /_____ _/ /_  / /__ 
 / __ `/ / / / __/ __ \    / __/ __ `/ __ \/ / _ \
/ /_/ / /_/ / /_/ /_/ /   / /_/ /_/ / /_/ / /  __/          
\__,_/\__,_/\__/\____/____\__/\__,_/_.___/_/\___/ 
                    /_____/                                                                                                                                                              
 auto_table v1.0
 By: J.Rosales

 ----------------------------------------------------------------------------------------------   
    '''
    print(banner_text)

banner()
# Get the file paths from the user
excel_file_path = input('\nEnter the path of the Excel file: ').strip("'").strip('"')
sh_name = input("\nPlease enter the name of the sheet: ")
word_file_path = input('\nEnter the path and how do you want to save the .docx: ').strip("'").strip('"')

root, ext = os.path.splitext(word_file_path)
if ext != ".docx":
    word_file_path = root + ".docx"

# Open the Excel workbook and get the sheet
wb = openpyxl.load_workbook(excel_file_path)
sheet = wb[sh_name]

# Create a new Word document
document = Document()

proceed = input("\nDo you want to proceed with LANDSCAPE LAYOUT for the Word Document(docx) ? (y/n): ")
if proceed.lower() == "y":
    # Landscape Layout code goes here
    section = document.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

# Add a table to the document and set the number of rows and columns
table = document.add_table(rows=sheet.max_row, cols=sheet.max_column)

# Iterate over the rows and columns of the Excel sheet and fill in the table cells
for i in range(1, sheet.max_row+1):
    for j in range(1, sheet.max_column+1):
        # Get the cell value
        value = str(sheet.cell(row=i, column=j).value)

        # Set the cell text
        cell = table.cell(i-1, j-1)
        cell.text = value

# remove double spaces
for row in table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = r.text.replace("  ", " ")


# Save the Word document
document.save(word_file_path)

# Open the Word document
word = win32com.client.Dispatch("Word.Application")
doc = word.Documents.Open(word_file_path)

# Select the table
table = doc.Tables(1)

# Iterate through the rows of the table
for row in range(1, table.Rows.Count + 1):
    # Select the cell in the "ID" column
    table.Cell(row, 3).Select()
    # Get the text of the cell
    text = word.Selection.Text
    # If the cell contains the letter C, color it dark red
    if "C" in text:
        word.Selection.Range.Shading.BackgroundPatternColor = 128
    # If the cell contains the letter H, color it red
    elif "H" in text:
        word.Selection.Range.Shading.BackgroundPatternColor = 255
    # If the cell contains the letter M, color it Orange
    elif "M" in text:
        word.Selection.Range.Shading.BackgroundPatternColor = 39423
    # If the cell contains the letter Y, color it yellow
    elif "L" in text:
        word.Selection.Range.Shading.BackgroundPatternColor = 65535

# Save and close the document
doc.Save()
doc.Close()
word.Quit()