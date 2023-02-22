import openpyxl
import os
import time
from docxtpl import DocxTemplate
from docx import Document
import win32com.client
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
import pandas as pd
from openpyxl.chart import (
    BarChart3D,
    Reference,
)

def banner():
    banner_text = '''
--------------------------------------------------------------------------------------------
 ______                   _______                                                        
(_____ \                 (_______)          _                          _                 
 _____) ) _____           _______  _   _  _| |_  ___   ____   _____  _| |_  _____   ____ 
|  __  / | ___ |         |  ___  || | | |(_   _)/ _ \ |    \ (____ |(_   _)| ___ | / ___)
| |  \ \ | ____| _______ | |   | || |_| |  | |_| |_| || | | |/ ___ |  | |_ | ____|| |    
|_|   |_||_____)(_______)|_|   |_||____/    \__)\___/ |_|_|_|\_____|   \__)|_____)|_|    
                                                                                                                                                

                                Re_Automater v1.0
                                By: J.Rosales

 ----------------------------------------------------------------------------------------------   
    '''
    print(banner_text)

banner()
time.sleep(3)
excel_file_path = input('\nEnter the path of the Excel file: ').strip("'").strip('"')
sh_name = input("\nPlease enter the name of the sheet: ")
temp_path = input('\nPlease enter the template .docx you want to use: ').strip("'").strip('"')
root, ext = os.path.splitext(temp_path)
if ext != ".docx":
    temp_path = root + ".docx"

# Escape Special Characters on xlsx file
wb = openpyxl.load_workbook(excel_file_path)
sheet = wb[sh_name]

for row in sheet.iter_rows():
    for cell in row:
        if isinstance(cell.value, str):
            cell.value = cell.value.replace('>', '&gt;')
            cell.value = cell.value.replace('<', '&lt;')
            cell.value = cell.value.replace('<=', '&lt;=')
            cell.value = cell.value.replace('>=', '&gt;=')

wb.save(excel_file_path)

 ### Create Pivot Table as png

df = pd.read_excel(excel_file_path, sheet_name=sh_name)
pivot_table = pd.pivot_table(df, values='Item', index='Risk', aggfunc='count')

# Reorder the index of the pivot table
pivot_table = pivot_table.reindex(['Critical', 'High', 'Medium', 'Low'])

# Create chart
fig, ax = plt.subplots()
pivot_table.plot.bar(title='Vulnerabilities x Risk', legend=False, grid=True, ax=ax)
plt.xlabel('Risk')
plt.ylabel('Item of Vulnerabilities')

# Color each bar separately
colors = {'Critical': 'darkred', 'High': 'red', 'Medium': 'orange', 'Low': 'yellow'}
for i, risk in enumerate(pivot_table.index):
    ax.bar(i, pivot_table.loc[risk, 'Item'], color=colors[risk])

# Add data labels to chart
for i, risk in enumerate(pivot_table.index):
    ax.text(i, pivot_table.loc[risk, 'Item'], str(pivot_table.loc[risk, 'Item']), ha='center', va='bottom')

# Change grid style
ax.grid(color='gray', linestyle=':', linewidth=0.3)

# Create table
table = ax.table(cellText=pivot_table.values, 
                 colLabels=pivot_table.columns, 
                 rowLabels=pivot_table.index,
                 cellLoc='center', 
                 loc='bottom',
                 bbox=[0, -0.55, 1, 0.25])
    

plt.savefig('chart.png', bbox_inches='tight', dpi=300)


# Load Template docx
template = DocxTemplate(temp_path)
image = InlineImage(template, 'chart.png')
table_conts = []

# Iterate over the rows and columns of the Excel sheet and fill in the table cells
for i in range(2, sheet.max_row+1):
    table_conts.append({
        'Index': i-1,
        'Item': sheet.cell(i, 1).value,
        'ID': sheet.cell(i, 2).value,
        'Risk': sheet.cell(i, 3).value,
        'Name': sheet.cell(i, 4).value,
        'CVE': sheet.cell(i, 5).value,
        'System': sheet.cell(i, 6).value,
        'Description': sheet.cell(i, 7).value,
        'CVSS_Base_Score': sheet.cell(i, 8).value,
        'CVSS_Temp_Score': sheet.cell(i, 9).value,
        'Solution': sheet.cell(i, 10).value,
    })

ask_client = input('\nEnter the name of the CLIENT: ')
ask_infra = input('\nEnter name of the INFRASTRUCTURE: ')
start_date = input('\nEnter Start Date: ')
end_date = input('\nEnter End Date: ')
context = {
    'NOMEINFRASTRUTTURA': ask_infra,
    'NOMECLIENTE': ask_client,
    'data_inizio': start_date,
    'data_fine': end_date,
    'table_contents': table_conts,
    'CHART': image,
}
template.render(context)
ask_fin = input('\nSelect the name of the final .docx output: ').strip("'").strip('"')
root, ext = os.path.splitext(ask_fin)
if ext != ".docx":
    ask_fin = root + ".docx"
time.sleep(2)
template.save(ask_fin)
time.sleep(2)
############# Edit Layout #######################

# Load the document
time.sleep(3)
docu = Document(ask_fin)

# Get the 6th table
table = docu.tables[5]

# Change the table to landscape orientation
proceed = input("\nDo you want to proceed with LANDSCAPE LAYOUT for the whole Word Document(docx) ? (y/n): ")
if proceed.lower() == "y":
    print('\nNow we are going to iterate with the Layout of the page [...]')
    section = docu.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

# remove double spaces
for row in table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = r.text.replace("  ", " ")
    
# Delete the 1st Column(Index)
def delete_columns(table, columns):
    # sort columns descending
    columns.sort(reverse=True)
    
    grid = table._tbl.find("w:tblGrid", table._tbl.nsmap)
    for ci in columns:
        for cell in table.column_cells(ci):
            cell._tc.getparent().remove(cell._tc)

        # Delete column reference.
        col_elem = grid[ci]
        grid.remove(col_elem)

# Delete columns 1 
delete_columns(table, [0])

# Save the modified document
time.sleep(2)
docu.save(ask_fin)

############# Color the Cell of Risk column #################

# Open the Word document
word = win32com.client.Dispatch("Word.Application")
doc = word.Documents.Open(ask_fin)

# Select the table
table = doc.Tables(6)

# Iterate through the rows of the table
for row in range(1, table.Rows.Count + 1):
    # Select the cell in the "Risk" column
    table.Cell(row, 3).Select()
    # Get the text of the cell
    text = word.Selection.Text
    # If the cell contains the letter C, color it dark red
    if "C" in text:
        word.Selection.Range.Shading.BackgroundPatternColor = 128
    # If the cell contains the letter H, color it red
    elif "H" in text:
        word.Selection.Range.Shading.BackgroundPatternColor = 255
    # If the cell contains the letter M, color it Orange
    elif "M" in text:
        word.Selection.Range.Shading.BackgroundPatternColor = 39423
    # If the cell contains the letter Y, color it yellow
    elif "L" in text:
        word.Selection.Range.Shading.BackgroundPatternColor = 65535

# Save and close the document
doc.Save()
doc.Close()
word.Quit()
time.sleep(3)

##### Generating Pivot Table in a new Excel file between Item and Risk ##########

print('\nGenerating Pivot Table in a new xlsx file [....]')

# Load data from excel file
time.sleep(3)
df = pd.read_excel(excel_file_path, sheet_name=sh_name)

# Create pivot table
pivot_table = pd.pivot_table(df, values='Item', index='Risk', aggfunc='count')
order = ['Critical', 'High', 'Medium', 'Low']
pivot_table = pivot_table.reindex(order)


# Save pivot table to excel file
ch_place = input('\nChoose the path and the name of the new .xlsx: ').strip("'").strip('"')
root, ext = os.path.splitext(ch_place)
if ext != ".xlsx":
    ch_place = root + ".xlsx"

writer = pd.ExcelWriter(ch_place, engine='openpyxl')
pivot_table.to_excel(writer, sheet_name='PivotChart')

# Create chart
book = writer.book
sheet = writer.sheets['PivotChart']

data = Reference(sheet, min_col=2, min_row=1, max_col=2, max_row=len(pivot_table)+1)
categories = Reference(sheet, min_col=1, min_row=2, max_row=len(pivot_table)+1, max_col=1)

chart = BarChart3D()
chart.title = "Vulnerabilities x Risk"
chart.add_data(data, titles_from_data=True)
chart.set_categories(categories)
chart.has_table = True

sheet.add_chart(chart, "E2")
writer.save() 